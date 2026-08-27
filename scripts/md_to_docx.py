"""将《EEG 设备接入指南.md》转换为富文本 Word 文档

支持的 markdown 元素：
- 一级/二级/三级/四级标题（H1-H4）
- 段落文本（含加粗、行内代码、链接）
- 有序列表 / 无序列表
- 表格（含表头）
- 代码块（等宽字体 + 浅灰背景）
- 引用块（左边框 + 灰色背景）
- 水平分隔线
- 图片嵌入 ![alt](path)
- ASCII 架构图（等宽字体居中显示）
"""

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches


# ============================================================
# 颜色与样式常量
# ============================================================

COLOR_H1 = RGBColor(0x7C, 0x3A, 0xED)       # 紫色
COLOR_H2 = RGBColor(0x5B, 0x21, 0xB6)       # 深紫
COLOR_H3 = RGBColor(0x1E, 0x40, 0xAF)       # 深蓝
COLOR_H4 = RGBColor(0x37, 0x41, 0x51)       # 深灰蓝
COLOR_CODE = RGBColor(0xBE, 0x18, 0x5D)     # 玫红（行内代码）
COLOR_LINK = RGBColor(0x1E, 0x40, 0xAF)     # 蓝色（链接）
COLOR_QUOTE = RGBColor(0x64, 0x74, 0x8B)    # 灰色（引用）
COLOR_TABLE_HEADER_BG = "7C3AED"            # 紫色表头背景
COLOR_TABLE_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_CODE_BG = "F3F4F6"                    # 浅灰代码块背景
COLOR_QUOTE_BG = "F5F3FF"                   # 浅紫引用背景


# ============================================================
# 辅助函数
# ============================================================

def set_cell_background(cell, color_hex: str):
    """设置表格单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_paragraph_background(paragraph, color_hex: str):
    """设置段落背景色（用于代码块和引用块）"""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    p_pr.append(shd)


def set_paragraph_border(paragraph, left: bool = False, color: str = "7C3AED", size: str = "18"):
    """设置段落边框（用于引用块的左边框）"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    if left:
        left_bdr = OxmlElement("w:left")
        left_bdr.set(qn("w:val"), "single")
        left_bdr.set(qn("w:sz"), size)
        left_bdr.set(qn("w:space"), "8")
        left_bdr.set(qn("w:color"), color)
        p_bdr.append(left_bdr)
    p_pr.append(p_bdr)


def add_inline_formatting(paragraph, text: str):
    """处理行内格式：加粗 **text**、行内代码 `code`、链接 [text](url)"""
    # 正则匹配：**bold** / `code` / [text](url)
    pattern = re.compile(
        r'(\*\*[^*]+\*\*)|(`[^`]+`)|(\[[^\]]+\]\([^)]+\))'
    )

    pos = 0
    for match in pattern.finditer(text):
        # 添加匹配前的普通文本
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.font.size = Pt(10.5)

        token = match.group()

        if token.startswith("**") and token.endswith("**"):
            # 加粗
            run = paragraph.add_run(token[2:-2])
            run.font.bold = True
            run.font.size = Pt(10.5)

        elif token.startswith("`") and token.endswith("`"):
            # 行内代码
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_CODE
            # 设置中文字体
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")

        elif token.startswith("[") and "](" in token:
            # 链接 [text](url)
            link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', token)
            if link_match:
                link_text = link_match.group(1)
                link_url = link_match.group(2)
                # 锚点链接（#xxx）只显示文字，不添加超链接
                if link_url.startswith("#"):
                    run = paragraph.add_run(link_text)
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = COLOR_LINK
                    run.font.underline = False
                else:
                    # 外部链接添加超链接
                    _add_hyperlink(paragraph, link_url, link_text)

        pos = match.end()

    # 添加剩余文本
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(10.5)


def _add_hyperlink(paragraph, url: str, text: str):
    """添加真正的超链接（可点击）"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # 字体设置
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    rPr.append(rFonts)

    # 蓝色 + 下划线
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1E40AF")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_font_for_run(run, font_name: str = "微软雅黑", size: int = 10.5, bold: bool = False,
                     color: RGBColor = None):
    """统一设置 run 的字体（中西文）"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


# ============================================================
# Markdown 解析与转换
# ============================================================

def is_table_separator(line: str) -> bool:
    """判断是否是表格分隔行 |---|---|"""
    return bool(re.match(r'^\|[\s\-:|]+\|$', line.strip()))


def parse_table_row(line: str) -> list[str]:
    """解析表格行，返回单元格列表"""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_ascii_diagram(line: str) -> bool:
    """判断是否是 ASCII 架构图行"""
    return bool(re.match(r'^[┌┐└┘│─├┤┬┴┼▼▲◀▶►◄◥◣◢█▓▒░ ]+$', line.strip())) or line.startswith("┌")


def convert_md_to_docx(md_path: str, docx_path: str):
    """将 markdown 文件转换为 Word 文档"""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0
    in_code_block = False
    code_buffer = []
    code_lang = ""

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # ---- 代码块处理 ----
        if line.strip().startswith("```"):
            if not in_code_block:
                # 开始代码块
                in_code_block = True
                code_lang = line.strip()[3:].strip()
                code_buffer = []
            else:
                # 结束代码块，输出
                in_code_block = False
                code_text = "\n".join(code_buffer)

                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.left_indent = Cm(0.5)
                    set_paragraph_background(p, COLOR_CODE_BG)

                    # 代码块标题（语言标识）
                    if code_lang:
                        lang_run = p.add_run(f"  [{code_lang}]")
                        set_font_for_run(lang_run, "Consolas", 9, False, RGBColor(0x9C, 0xA3, 0xAF))
                        p.add_run("\n")

                    # 代码内容
                    code_run = p.add_run(code_text)
                    set_font_for_run(code_run, "Consolas", 9.5, False, RGBColor(0x1F, 0x29, 0x37))
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # ---- 空行 ----
        if not line.strip():
            i += 1
            continue

        # ---- 水平分隔线 ----
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # 添加底部边框作为分隔线
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "D1D5DB")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
            i += 1
            continue

        # ---- 图片 ----
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', line)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2).strip()

            # 解析相对路径（相对于 md 文件所在目录的上级）
            md_dir = Path(md_path).resolve().parent
            if not os.path.isabs(img_path):
                # docs/ 目录的上级是项目根目录
                img_full_path = (md_dir / img_path).resolve()
            else:
                img_full_path = img_path

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)

            if os.path.exists(img_full_path):
                try:
                    run = p.add_run()
                    # 限制图片宽度为页面可用宽度
                    run.add_picture(str(img_full_path), width=Inches(6.0))
                    print(f"  ✅ 插入图片：{img_path}")
                except Exception as e:
                    run = p.add_run(f"[图片插入失败: {img_path} - {e}]")
                    set_font_for_run(run, "微软雅黑", 9, False, RGBColor(0xDC, 0x26, 0x26))
            else:
                run = p.add_run(f"[图片未找到: {img_path}]")
                set_font_for_run(run, "微软雅黑", 9, False, RGBColor(0xDC, 0x26, 0x26))

            # 图片说明（alt 文本作为图注）
            if alt_text:
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(8)
                cap_run = caption.add_run(f"图：{alt_text}")
                set_font_for_run(cap_run, "微软雅黑", 9, False, RGBColor(0x6B, 0x72, 0x80))
                cap_run.font.italic = True

            i += 1
            continue

        # ---- 标题 ----
        title_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if title_match:
            level = len(title_match.group(1))
            title_text = title_match.group(2).strip()

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
            p.paragraph_format.space_after = Pt(6)

            run = p.add_run(title_text)
            if level == 1:
                set_font_for_run(run, "微软雅黑", 22, True, COLOR_H1)
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(12)
            elif level == 2:
                set_font_for_run(run, "微软雅黑", 16, True, COLOR_H2)
            elif level == 3:
                set_font_for_run(run, "微软雅黑", 13, True, COLOR_H3)
            else:
                set_font_for_run(run, "微软雅黑", 11.5, True, COLOR_H4)

            i += 1
            continue

        # ---- 引用块 ----
        if line.startswith(">"):
            quote_text = line.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            set_paragraph_background(p, COLOR_QUOTE_BG)
            set_paragraph_border(p, left=True, color="7C3AED", size="18")
            add_inline_formatting(p, quote_text)
            # 引用块文字颜色调灰
            for run in p.runs:
                if not run.font.color.rgb or run.font.color.rgb == RGBColor(0, 0, 0):
                    run.font.color.rgb = COLOR_QUOTE
            i += 1
            continue

        # ---- 表格 ----
        if line.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            # 解析表格
            header = parse_table_row(line)
            i += 2  # 跳过分隔行
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1

            # 创建 Word 表格
            n_cols = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=n_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"

            # 表头
            header_cells = table.rows[0].cells
            for j, cell_text in enumerate(header):
                if j < n_cols:
                    cell = header_cells[j]
                    set_cell_background(cell, COLOR_TABLE_HEADER_BG)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(cell_text)
                    set_font_for_run(run, "微软雅黑", 10, True, COLOR_TABLE_HEADER_TEXT)

            # 数据行
            for row_idx, row_data in enumerate(rows):
                row_cells = table.rows[row_idx + 1].cells
                for j, cell_text in enumerate(row_data):
                    if j < n_cols:
                        cell = row_cells[j]
                        # 隔行变色
                        if row_idx % 2 == 1:
                            set_cell_background(cell, "F9FAFB")
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        add_inline_formatting(p, cell_text)

            # 表格后空一行
            doc.add_paragraph()
            continue

        # ---- 有序列表 ----
        ol_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if ol_match:
            content = ol_match.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.8)
            add_inline_formatting(p, content)
            i += 1
            continue

        # ---- 无序列表 ----
        ul_match = re.match(r'^[-*]\s+(.+)$', line)
        if ul_match:
            content = ul_match.group(1)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.8)
            add_inline_formatting(p, content)
            i += 1
            continue

        # ---- ASCII 架构图 ----
        if is_ascii_diagram(line):
            # 收集连续的 ASCII 图行
            diagram_lines = [line]
            i += 1
            while i < len(lines) and (is_ascii_diagram(lines[i].rstrip("\n")) or
                                      (lines[i].strip() and not lines[i].startswith("#") and
                                       not lines[i].startswith("|") and
                                       not lines[i].startswith("-") and
                                       not lines[i].startswith(">") and
                                       not lines[i].startswith("```") and
                                       not re.match(r'^\d+\.', lines[i]) and
                                       not re.match(r'^[-*]\s', lines[i]))):
                diagram_lines.append(lines[i].rstrip("\n"))
                i += 1

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            set_paragraph_background(p, COLOR_CODE_BG)

            diagram_text = "\n".join(diagram_lines)
            run = p.add_run(diagram_text)
            set_font_for_run(run, "Consolas", 8.5, False, RGBColor(0x1F, 0x29, 0x37))
            continue

        # ---- 普通段落 ----
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        add_inline_formatting(p, line)
        i += 1

    # 保存文档
    doc.save(docx_path)
    print(f"✅ Word 文档已生成：{docx_path}")
    print(f"   文件大小：{os.path.getsize(docx_path) / 1024:.1f} KB")


# ============================================================
# 主入口
# ============================================================

if __name__ == "__main__":
    ROOT = Path(__file__).resolve().parent.parent
    md_file = ROOT / "docs" / "EEG设备接入指南.md"
    docx_file = ROOT / "docs" / "EEG设备接入指南.docx"

    if not md_file.exists():
        print(f"✗ 源文件不存在：{md_file}")
        sys.exit(1)

    convert_md_to_docx(str(md_file), str(docx_file))
