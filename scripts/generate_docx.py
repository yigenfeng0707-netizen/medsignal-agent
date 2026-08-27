"""生成项目申报书 Word 文档（高质量富文本版）"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import os

# ============ 配色方案 ============
COLOR_PRIMARY = RGBColor(0x1A, 0x56, 0xDB)      # 主色：科技蓝
COLOR_SECONDARY = RGBColor(0x0E, 0x7C, 0x86)    # 副色：医疗青
COLOR_ACCENT = RGBColor(0xDC, 0x26, 0x26)       # 强调：警示红
COLOR_DARK = RGBColor(0x1F, 0x29, 0x37)         # 深色文字
COLOR_GRAY = RGBColor(0x6B, 0x72, 0x80)         # 灰色辅助
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_LIGHT_BG = RGBColor(0xF0, 0xF5, 0xFF)     # 浅蓝背景
COLOR_HIGHLIGHT = RGBColor(0xFE, 0xF3, 0xC7)    # 高亮黄
COLOR_STAR = RGBColor(0xF5, 0x9E, 0x0B)         # 星标橙

# ============ 工具函数 ============
def set_cell_bg(cell, color_hex):
    """设置单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            border.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            border.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tc_borders.append(border)
    tc_pr.append(tc_borders)

def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()

def set_paragraph_shading(paragraph, color_hex):
    """设置段落底纹"""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    p_pr.append(shd)

def add_horizontal_line(paragraph, color="1A56DB"):
    """在段落底部添加水平线"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

# ============ 文档样式设置 ============
def setup_styles(doc):
    """设置文档默认样式"""
    # 正文样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_DARK
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # 标题1样式
    h1 = doc.styles['Heading 1']
    h1.font.name = '微软雅黑'
    h1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_PRIMARY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)

    # 标题2样式
    h2 = doc.styles['Heading 2']
    h2.font.name = '微软雅黑'
    h2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_SECONDARY
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)

    # 标题3样式
    h3 = doc.styles['Heading 3']
    h3.font.name = '微软雅黑'
    h3.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_DARK
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(6)

# ============ 内容构建函数 ============
def add_cover(doc):
    """封面页"""
    # 顶部留白
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('项目申报书')
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('全球脑机接口×医保创新场景大赛')
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_SECONDARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('场景应用赛道')
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_GRAY

    # 装饰线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 20)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.size = Pt(14)

    # 项目名称
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('医保智脑')
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('基于BCI脑电监测的医保主动健康服务智能体')
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_GRAY

    # 底部信息
    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026年7月')
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_GRAY

    add_page_break(doc)


def add_section_title(doc, number, title):
    """添加章节标题（带编号和装饰线）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f'{number}  {title}')
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    add_horizontal_line(p, "1A56DB")


def add_subsection_title(doc, number, title):
    """添加二级标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{number}  {title}')
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY


def add_body_text(doc, text, bold=False, color=None, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.color.rgb = color or COLOR_DARK
    if bold:
        run.font.bold = True
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    """添加项目符号段落"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.74 + level * 0.5)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        run2 = p.add_run(text)
        run2.font.name = '微软雅黑'
        run2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(11)
        run2.font.color.rgb = COLOR_DARK
    else:
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_DARK


def add_numbered(doc, text, bold_prefix=None):
    """添加编号段落"""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        run2 = p.add_run(text)
        run2.font.name = '微软雅黑'
        run2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(11)
        run2.font.color.rgb = COLOR_DARK
    else:
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_DARK


def add_callout_box(doc, title, content, color="FEF3C7", border_color="F59E0B"):
    """添加提示框（带背景色和边框）"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, color)
    set_cell_border(cell,
        top={'sz': 4, 'color': border_color},
        bottom={'sz': 4, 'color': border_color},
        left={'sz': 12, 'color': border_color},
        right={'sz': 4, 'color': border_color}
    )
    # 标题
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x9A, 0x34, 0x12)
    # 内容
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(content)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_DARK
    doc.add_paragraph()


def add_styled_table(doc, headers, rows, col_widths=None, highlight_rows=None):
    """添加专业样式表格（带表头配色、斑马纹、高亮行）"""
    highlight_rows = highlight_rows or []
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        set_cell_bg(cell, "1A56DB")
        set_cell_border(cell,
            top={'sz': 8, 'color': '1A56DB'},
            bottom={'sz': 8, 'color': '1A56DB'},
            left={'sz': 4, 'color': '1A56DB'},
            right={'sz': 4, 'color': '1A56DB'}
        )
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(header)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_WHITE

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        is_highlight = row_idx in highlight_rows
        for col_idx, cell_data in enumerate(row_data):
            cell = row_cells[col_idx]
            if is_highlight:
                set_cell_bg(cell, "FEF3C7")
            elif row_idx % 2 == 1:
                set_cell_bg(cell, "F0F5FF")
            set_cell_border(cell,
                top={'sz': 2, 'color': 'D1D5DB'},
                bottom={'sz': 2, 'color': 'D1D5DB'},
                left={'sz': 2, 'color': 'D1D5DB'},
                right={'sz': 2, 'color': 'D1D5DB'}
            )
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            # 处理加粗标记
            text = str(cell_data)
            run = p.add_run(text)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(10.5)
            run.font.color.rgb = COLOR_DARK
            if is_highlight:
                run.font.bold = True
    doc.add_paragraph()


def add_info_card(doc, label, value, label_color=COLOR_PRIMARY):
    """添加信息卡片（标签:值）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(label)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = label_color
    run2 = p.add_run(value)
    run2.font.name = '微软雅黑'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run2.font.size = Pt(11)
    run2.font.color.rgb = COLOR_DARK


# ============ 主文档生成 ============
def generate_docx():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    setup_styles(doc)

    # ===== 封面 =====
    add_cover(doc)

    # ===== 一、项目基本信息 =====
    add_section_title(doc, '一、', '项目基本信息')

    add_info_card(doc, '项目名称：', '医保智脑——基于BCI脑电监测的医保主动健康服务智能体')
    add_info_card(doc, '参赛赛道：', '场景应用赛道（赛道7：芯测预警——脑血管疾病预警、认知与精神状态早期筛查、行业人群脑健康监测）', COLOR_SECONDARY)
    add_info_card(doc, '项目定位：', '将脑机接口（BCI）脑电监测技术与医保数据要素深度融合，构建"脑电采集→健康评估→风险预警→医保政策联动"的全链路智能服务，实现从"被动报销"到"主动健康"的范式转变。', COLOR_DARK)

    # ===== 二、项目背景与痛点 =====
    add_section_title(doc, '二、', '项目背景与痛点')

    add_subsection_title(doc, '2.1', '社会背景')
    add_body_text(doc, '我国脑血管疾病年发病人数超300万，认知障碍患者超5000万，焦虑抑郁人群超9000万。早期发现、早期干预可显著降低疾病负担。然而存在以下痛点：', indent=True)

    add_numbered(doc, '常规体检缺乏脑功能筛查，脑血管病首发即重症比例高', '筛查覆盖率低：')
    add_numbered(doc, '轻度认知障碍（MCI）早期无明显症状，确诊时已进展', '认知衰退隐匿：')
    add_numbered(doc, '焦虑抑郁早期识别率不足30%，错过最佳干预窗口', '精神健康忽视：')
    add_numbered(doc, '80%参保人不知道可享受的脑健康相关医保待遇', '医保政策信息不对称：')

    add_subsection_title(doc, '2.2', '技术机遇')
    add_body_text(doc, '脑机接口技术使脑电监测从实验室走向消费级场景。结合医保数据要素战略，可实现：', indent=True)
    add_bullet(doc, '自动识别脑血管/认知/精神风险', '脑电异常 → ')
    add_bullet(doc, '自动联动医保政策（门诊慢病/检查报销/绿色通道）', '风险预警 → ')
    add_bullet(doc, '从"生病报销"到"预警预防"', '主动健康 → ')

    # ===== 三、技术方案 =====
    add_section_title(doc, '三、', '技术方案')

    add_subsection_title(doc, '3.1', '系统架构')
    add_body_text(doc, '采用"1+6+N"多智能体架构：', indent=True)
    add_bullet(doc, '意图识别 + 路由分发 + 结果聚合', '1个编排智能体：')
    add_bullet(doc, '权益管家、报销助手、健康卫士、政策参谋、安全守门、脑电卫士（BCI创新）', '6个专业智能体：')
    add_bullet(doc, '权益查询、报销预审、健康画像、政策匹配、脑电健康评估', 'N个应用场景：')

    add_subsection_title(doc, '3.2', 'BCI脑电健康引擎（核心创新）')
    add_info_card(doc, '信号采集：', '支持消费级EEG设备（Muse 4通道 TP9/AF7/AF8/TP10，256Hz）+ 合成信号模拟 + 文件导入')
    add_info_card(doc, '频域分析：', 'Welch PSD（Hann窗+50%重叠+FFT）→ δ/θ/α/β/γ 五频段功率积分')

    add_body_text(doc, '健康指标计算（8维，0-100归一化）：', bold=True)
    add_styled_table(doc,
        headers=['指标', '算法依据', '临床意义'],
        rows=[
            ['压力指数', 'α/β比值反演', '长期压力影响心血管'],
            ['注意力指数', 'θ/β比值反演', '注意力下降影响生活'],
            ['睡眠质量', 'δ+θ占比', '深睡相关频段'],
            ['认知负荷', 'β+γ占比', '脑力消耗程度'],
            ['情绪状态', 'α不对称性+β活跃度', '焦虑/低落倾向'],
            ['脑血管风险 ⭐', 'δ波激增+θ/α升高+α波抑制', '脑血管功能异常早期预警'],
            ['认知衰退风险 ⭐', 'θ/α比值升高+α波相对功率降低', '轻度认知障碍(MCI)筛查'],
            ['精神状态筛查 ⭐', '焦虑评分+抑郁评分综合', '焦虑/抑郁倾向量化'],
        ],
        col_widths=[4, 6, 6],
        highlight_rows=[5, 6, 7]
    )

    add_body_text(doc, '异常预警（8条规则，带evidence可解释）：', bold=True)
    add_bullet(doc, '高压力/睡眠差/认知过载/注意力低/情绪异常', '常规5条：')
    add_bullet(doc, '脑血管风险≥60 / 认知衰退≥60 / 精神状态焦虑抑郁≥60', '赛道7核心3条：')

    add_body_text(doc, '医保政策联动（8条规则）：', bold=True)
    add_bullet(doc, '脑血管病门诊慢病待遇 / 脑卒中绿色通道 / 颈动脉超声医保报销')
    add_bullet(doc, '老年认知功能筛查 / 阿尔茨海默病门诊慢病 / 认知康复训练报销')
    add_bullet(doc, '心理治疗医保报销 / 抑郁症焦虑症门诊慢病 / PHQ-9/GAD-7量表筛查')

    add_subsection_title(doc, '3.3', '医保数据要素对齐')
    add_body_text(doc, '对齐浙江省医保数据要素战略"1+3+N"框架：', indent=True)
    add_bullet(doc, '统一数据治理', '1个数据底座：')
    add_bullet(doc, '隐私计算（可用不可见）、区块链存证（可信可追溯）、数据沙箱（可控可计量）', '3个核心能力：')
    add_bullet(doc, '脑电健康评估即为典型场景', 'N个应用场景：')

    add_subsection_title(doc, '3.4', '全链路可解释性')
    add_body_text(doc, '每个决策都有来源：脑电指标有频段依据，预警有阈值规则，政策推荐有原文引用。不做黑箱AI，做参保人能信任的AI。', indent=True)

    # ===== 四、创新亮点 =====
    add_section_title(doc, '四、', '创新亮点')
    add_numbered(doc, '国内首个将脑电监测与医保政策联动，实现"脑电异常→风险预警→医保政策自动推荐"', 'BCI×医保全链路创新：')
    add_numbered(doc, '脑血管预警+认知衰退筛查+精神状态筛查，均有临床脑电标志物依据', '赛道7三大能力全覆盖：')
    add_numbered(doc, '6个专业智能体+编排器，支持脑电+政策复合意图并行调度', '多智能体协作：')
    add_numbered(doc, '支持Muse等消费级设备，降低脑电监测门槛', '消费级BCI适配：')
    add_numbered(doc, '从"生病找医保"到"脑电预警+医保主动服务"', '主动健康范式：')

    # ===== 五、应用场景 =====
    add_section_title(doc, '五、', '应用场景')
    add_numbered(doc, '社区卫生服务中心配备EEG设备，老年人定期脑电筛查→脑血管/认知风险预警→医保政策指导', '社区脑健康筛查：')
    add_numbered(doc, '高危行业（交通/医疗/能源）员工定期脑电监测→疲劳/压力预警→心理健康医保服务', '职场脑健康监测：')
    add_numbered(doc, '糖尿病/高血压患者脑电监测→脑血管并发症早期预警→门诊慢病待遇联动', '慢病管理增强：')
    add_numbered(doc, '脑电精神状态筛查→焦虑抑郁倾向量化→心理治疗医保报销引导', '精神健康早期识别：')

    # ===== 六、产业化路径 =====
    add_section_title(doc, '六、', '产业化路径')

    add_styled_table(doc,
        headers=['阶段', '时间', '目标'],
        rows=[
            ['Phase 1', '1-3个月', '接入浙江省医保数据赋能实验室沙箱，脱敏数据验证'],
            ['Phase 2', '3-6个月', '选择1-2个区县试点，覆盖5-10万参保人'],
            ['Phase 3', '6-12个月', '全省推广，依托医保数据要素基础设施'],
        ],
        col_widths=[3, 3, 10]
    )

    add_body_text(doc, '商业模式：', bold=True)
    add_bullet(doc, '为医保部门提供智能客服+脑健康筛查服务，降低经办成本', 'To G：')
    add_bullet(doc, '为商保/药企提供脑健康数据赋能，支撑精准定价', 'To B：')
    add_bullet(doc, '基础脑电评估免费，深度健康管理增值服务付费', 'To C：')

    # ===== 七、团队信息 =====
    add_section_title(doc, '七、', '团队信息')

    add_callout_box(doc,
        '【需用户填写】',
        '以下信息需参赛者补充真实团队信息。注意：申报书中可填写团队信息，但路演材料中不得出现参赛者身份信息。',
        color="FEF3C7", border_color="F59E0B"
    )

    add_info_card(doc, '团队名称：', '_________________（请填写）')
    add_body_text(doc, '团队成员：', bold=True)
    add_numbered(doc, '负责 _____，技术栈 _____', '_____（负责人）：')
    add_numbered(doc, '负责 _____，技术栈 _____', '_____（成员）：')
    add_info_card(doc, '团队优势：', '_________________（请填写团队在BCI/医保/AI领域的能力与经验）')
    add_info_card(doc, '联系方式：', '_________________（请填写，仅用于大赛组委会联系，不会出现在路演材料中）')

    # ===== 八、项目成果 =====
    add_section_title(doc, '八、', '项目成果')

    add_styled_table(doc,
        headers=['维度', '数据'],
        rows=[
            ['后端API', '30+ 端点（7个Router + EEG模块）'],
            ['前端页面', '9个（首页/对话/权益/报销/健康/政策/安全/脑电/数据空间）'],
            ['算法引擎', '4大引擎（报销计算/健康画像/政策匹配/EEG脑电）'],
            ['智能体', '6个专业Agent + 1个编排器'],
            ['脑电指标', '8维（含赛道7三大核心）'],
            ['预警规则', '8条（含赛道7三条核心）'],
            ['政策联动', '8条（含赛道7三条核心）'],
            ['单元测试', '151项全绿'],
            ['冒烟测试', '68项全绿'],
            ['医保规则库', '浙江省/杭州市参数（3险种×3级别×门诊住院）'],
        ],
        col_widths=[4, 12]
    )

    # ===== 九、社会价值 =====
    add_section_title(doc, '九、', '社会价值')
    add_numbered(doc, '推动医保从"被动报销"到"主动健康"', '范式转变：')
    add_numbered(doc, '脑电预警让脑血管病/认知障碍/精神疾病早发现早干预', '早筛早防：')
    add_numbered(doc, '让5600万浙江参保人知道并享受到脑健康相关医保待遇', '政策普惠：')
    add_numbered(doc, 'BCI数据×医保数据的融合创新示范', '数据要素价值释放：')

    # ===== 项目承诺 =====
    add_section_title(doc, '十、', '项目承诺')
    add_callout_box(doc,
        '原创与合规承诺',
        '本项目为原创开发，所使用数据均为脱敏演示数据，不涉及真实参保人隐私。生产部署将严格遵守《个人信息保护法》《数据安全法》及医保数据管理相关规定。',
        color="F0F5FF", border_color="1A56DB"
    )

    # 保存
    output_path = r'd:\浙江医保数据应用赛道\yibao-eeg\docs\项目申报书.docx'
    doc.save(output_path)
    print(f'Word文档已生成：{output_path}')
    print(f'   文件大小：{os.path.getsize(output_path) / 1024:.1f} KB')
    return output_path


if __name__ == '__main__':
    generate_docx()
